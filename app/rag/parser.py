import csv
import json
import re
from collections.abc import Callable
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any


@dataclass(frozen=True)
class ParsedDocument:
    title: str
    content: str
    images: dict[str, str] = field(default_factory=dict)
    metadata: dict[str, Any] = field(default_factory=dict)
    pages: list[dict[str, Any]] = field(default_factory=list)

    @property
    def text(self) -> str:
        return self.content


ParserCallable = Callable[[Path], ParsedDocument]

MINERU_FILE_TYPES = [
    "pdf",
    "doc",
    "docx",
    "ppt",
    "pptx",
    "xls",
    "xlsx",
    "png",
    "jpg",
    "jpeg",
    "jp2",
    "webp",
    "gif",
    "bmp",
]


class ParserEngineRegistry:
    def __init__(self) -> None:
        self._engines: dict[str, dict[str, ParserCallable]] = {}
        self._descriptions: dict[str, str] = {}
        self._availability: dict[str, tuple[bool, str]] = {}
        self._register_builtin()
        self.register(
            "mineru",
            {},
            "MinerU OCR / advanced parser",
            available=True,
            advertised_file_types=MINERU_FILE_TYPES,
        )

    def register(
        self,
        name: str,
        file_types: dict[str, ParserCallable],
        description: str,
        *,
        available: bool = True,
        unavailable_reason: str = "",
        advertised_file_types: list[str] | None = None,
    ) -> None:
        handlers = {key.lower().lstrip("."): value for key, value in file_types.items()}
        if advertised_file_types:
            for file_type in advertised_file_types:
                handlers.setdefault(file_type.lower().lstrip("."), self._unsupported_parser)
        self._engines[name] = handlers
        self._descriptions[name] = description
        self._availability[name] = (available, unavailable_reason)

    def resolve(self, engine: str | None, file_type: str) -> ParserCallable:
        normalized = file_type.lower().lstrip(".")
        if engine and engine in self._engines:
            handler = self._engines[engine].get(normalized)
            if handler and self._availability[engine][0]:
                return handler
        handler = self._engines["builtin"].get(normalized)
        if handler:
            return handler
        raise ValueError(f"Unsupported file type: {normalized or '<none>'}")

    def list_engines(self) -> list[dict[str, Any]]:
        engines = []
        for name, handlers in self._engines.items():
            available, reason = self._availability[name]
            engines.append(
                {
                    "name": name,
                    "description": self._descriptions[name],
                    "file_types": sorted(handlers),
                    "available": available,
                    "unavailable_reason": reason,
                }
            )
        return engines

    def _register_builtin(self) -> None:
        self.register(
            "builtin",
            {
                "txt": self._parse_text,
                "md": self._parse_markdown,
                "markdown": self._parse_markdown,
                "pdf": self._parse_pdf,
                "docx": self._parse_docx,
                "csv": self._parse_csv,
                "json": self._parse_json,
                "xlsx": self._parse_xlsx,
            },
            "内置解析引擎",
        )

    def _metadata(self, path: Path) -> dict[str, Any]:
        return {"file_name": path.name, "file_type": path.suffix.lower().lstrip(".")}

    def _parse_text(self, path: Path) -> ParsedDocument:
        return ParsedDocument(
            title=path.name,
            content=path.read_text(encoding="utf-8").strip(),
            metadata=self._metadata(path),
        )

    def _parse_markdown(self, path: Path) -> ParsedDocument:
        content = path.read_text(encoding="utf-8").strip()
        return ParsedDocument(title=path.name, content=_format_markdown_tables(content), metadata=self._metadata(path))

    def _parse_pdf(self, path: Path) -> ParsedDocument:
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        page_texts = [page.extract_text() or "" for page in reader.pages]
        pages = [
            {"page": index + 1, "start": 0, "end": len(text)}
            for index, text in enumerate(page_texts)
        ]
        metadata = {**self._metadata(path), "page_count": len(page_texts)}
        return ParsedDocument(title=path.name, content="\f".join(page_texts).strip(), metadata=metadata, pages=pages)

    def _parse_docx(self, path: Path) -> ParsedDocument:
        from docx import Document

        doc = Document(str(path))
        lines: list[str] = []
        lines.extend(paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip())
        for table in doc.tables:
            rows = [[" ".join(cell.text.split()) for cell in row.cells] for row in table.rows]
            if rows:
                lines.append(_rows_to_markdown(rows))
        return ParsedDocument(title=path.name, content="\n\n".join(lines).strip(), metadata=self._metadata(path))

    def _parse_csv(self, path: Path) -> ParsedDocument:
        with path.open("r", encoding="utf-8-sig", newline="") as file:
            rows = list(csv.reader(file))
        return ParsedDocument(title=path.name, content=_rows_to_markdown(rows), metadata=self._metadata(path))

    def _parse_json(self, path: Path) -> ParsedDocument:
        data = json.loads(path.read_text(encoding="utf-8"))
        content = "```json\n" + json.dumps(data, ensure_ascii=False, indent=2) + "\n```"
        return ParsedDocument(title=path.name, content=content, metadata=self._metadata(path))

    def _parse_xlsx(self, path: Path) -> ParsedDocument:
        from openpyxl import load_workbook

        workbook = load_workbook(path, read_only=True, data_only=True)
        sections: list[str] = []
        for sheet in workbook.worksheets:
            rows = [
                ["" if cell is None else str(cell) for cell in row]
                for row in sheet.iter_rows(values_only=True)
                if any(cell is not None and str(cell).strip() for cell in row)
            ]
            if rows:
                sections.append(f"## {sheet.title}\n\n{_rows_to_markdown(rows)}")
        return ParsedDocument(title=path.name, content="\n\n".join(sections), metadata=self._metadata(path))

    def _unsupported_parser(self, path: Path) -> ParsedDocument:
        raise ValueError(f"Unsupported file type: {path.suffix.lower().lstrip('.') or '<none>'}")


class DocumentParser:
    def __init__(self, registry: ParserEngineRegistry | None = None) -> None:
        self.registry = registry or ParserEngineRegistry()

    def parse(self, path: Path, engine: str | None = None) -> ParsedDocument:
        file_type = path.suffix.lower().lstrip(".")
        parser = self.registry.resolve(engine, file_type)
        return parser(path)


def _format_markdown_tables(content: str) -> str:
    lines = content.splitlines()
    out: list[str] = []
    for line in lines:
        stripped = line.strip()
        if stripped.startswith("|") and stripped.endswith("|"):
            cells = [cell.strip() for cell in stripped.strip("|").split("|")]
            if cells and all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in cells):
                cells = [_normalize_alignment(cell) for cell in cells]
            out.append("| " + " | ".join(cells) + " |")
        else:
            out.append(line)
    return "\n".join(out)


def _normalize_alignment(value: str) -> str:
    compact = value.replace(" ", "")
    left = ":" if compact.startswith(":") else ""
    right = ":" if compact.endswith(":") else ""
    return f"{left}---{right}"


def _rows_to_markdown(rows: list[list[str]]) -> str:
    if not rows:
        return ""
    width = max(len(row) for row in rows)
    normalized = [row + [""] * (width - len(row)) for row in rows]
    header = "| " + " | ".join(normalized[0]) + " |"
    sep = "| " + " | ".join("---" for _ in range(width)) + " |"
    body = ["| " + " | ".join(row) + " |" for row in normalized[1:]]
    return "\n".join([header, sep, *body])
